#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
使用方法：
1. 在下方 SOURCE_FILE 中填写要转换的文件路径
2. 运行:python toword.py
"""
import os
from datetime import datetime
from pathlib import Path

# 关键修复：导入python-docx的必要模块
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# ============================================================

# 要转换的文件路径（修改这里）
SOURCE_FILE = r"C:\Users\purriste\Desktop\PYProject\rppg\frontend\pages\health.vue"

# 输出目录（Word 文档保存位置）
OUTPUT_DIR = r"C:\Users\purriste\Desktop\PYProject\rppg"

# ============================================================


def get_file_language(file_extension: str) -> str:
    """根据文件扩展名获取语言标识"""
    language_map = {
        '.py': 'Python',
        '.vue': 'Vue',
        '.js': 'JavaScript',
        '.ts': 'TypeScript',
        '.html': 'HTML',
        '.css': 'CSS',
        '.json': 'JSON',
    }
    return language_map.get(file_extension.lower(), 'Code')


def convert_to_word(source_file: str, output_dir: str) -> bool:
    """
    将文件转换为 Word 文档
    
    Args:
        source_file: 源文件路径
        output_dir: 输出目录路径
    
    Returns:
        转换是否成功
    """
    source_path = Path(source_file).resolve()  # 修复：获取绝对路径，避免相对路径问题
    output_path = Path(output_dir).resolve()
    
    # 检查文件是否存在
    if not source_path.exists():
        print(f"错误: 文件不存在 - {source_path}")
        return False
    
    # 修复：添加文件读取的异常处理
    try:
        # 读取源文件内容（增加编码容错）
        with open(source_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        # 尝试其他常见编码
        try:
            with open(source_path, 'r', encoding='gbk') as f:
                content = f.read()
            print(f"提示: 文件 {source_path.name} 使用 GBK 编码读取")
        except Exception as e:
            print(f"错误: 读取文件失败 - {e}")
            return False
    except Exception as e:
        print(f"错误: 读取文件失败 - {e}")
        return False
    
    # 获取文件信息
    file_name = source_path.name
    file_extension = source_path.suffix
    language = get_file_language(file_extension)
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading(file_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加文件信息
    info_para = doc.add_paragraph()
    info_para.add_run("源文件路径: ").bold = True
    info_para.add_run(str(source_path))
    
    info_para2 = doc.add_paragraph()
    info_para2.add_run("文件类型: ").bold = True
    info_para2.add_run(language)
    
    info_para3 = doc.add_paragraph()
    info_para3.add_run("转换时间: ").bold = True
    info_para3.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    
    # 添加分隔线
    doc.add_paragraph("─" * 50)
    
    # 添加源代码标题
    doc.add_heading("源代码", level=1)
    
    # 添加代码内容（使用等宽字体）
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(content)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(10)
    
    # 确保输出目录存在
    output_path.mkdir(parents=True, exist_ok=True)
    
    # 生成输出文件名
    output_file = output_path / f"{source_path.stem}.docx"
    
    # 如果文件已存在，添加时间戳
    if output_file.exists():
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_file = output_path / f"{source_path.stem}_{timestamp}.docx"
    
    # 修复：添加保存文档的异常处理
    try:
        doc.save(output_file)
        print(f"\n[OK] 转换成功!")
        print(f"     源文件: {source_path}")
        print(f"     输出: {output_file}")
        return True
    except Exception as e:
        print(f"\n[ERROR] 保存 Word 文档失败 - {e}")
        return False


def main():
    """主函数"""
    print("=" * 60)
    print("文件转 Word 文档工具")
    print("=" * 60)
    
    print(f"\n源文件: {SOURCE_FILE}")
    print(f"输出目录: {OUTPUT_DIR}")
    
    # 执行转换并处理返回结果
    success = convert_to_word(SOURCE_FILE, OUTPUT_DIR)
    if not success:
        print("\n[ERROR] 转换失败！")
        exit(1)  # 非0退出码表示执行失败


if __name__ == '__main__':
    # 修复：添加整体异常捕获
    try:
        main()
    except Exception as e:
        print(f"\n[FATAL] 程序运行出错 - {e}")
        exit(1)